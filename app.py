#!/usr/bin/env python3
"""
Tab Scraper & Formatter
Scrape guitar tabs and output clean text / docx / gp files.

Usage:
    python app.py "https://tabs.ultimate-guitar.com/tab/..."
    python app.py --clean
"""

import sys
import re
import json
import argparse
import os

try:
    from curl_cffi import requests
    from bs4 import BeautifulSoup
except ImportError:
    print("Install dependencies:")
    print("  pip install curl_cffi beautifulsoup4")
    sys.exit(1)


# ─── Nettoyage texte ───────────────────────────────────────────────────────────

def normalize_whitespace(text: str) -> str:
    for ch in [' ', ' ', ' ']:
        text = text.replace(ch, ' ')
    for ch in ['​', '‌', '‍', '﻿']:
        text = text.replace(ch, '')
    return text


def fix_tab_lines(text: str) -> str:
    lines = text.split('\n')
    result = []
    tab_buffer = {}
    tab_order = ['e', 'B', 'G', 'D', 'A', 'E', 'd', 'b', 'g']

    def flush():
        out = [tab_buffer[s] for s in tab_order if s in tab_buffer]
        tab_buffer.clear()
        return out

    for line in lines:
        stripped = line.rstrip()
        m = re.match(r'^([eEBbGgDdAa])\|(.*)$', stripped)
        if m:
            note = m.group(1)
            if note in tab_buffer:
                result.extend(flush())
            tab_buffer[note] = stripped
        else:
            result.extend(flush())
            result.append(stripped)

    result.extend(flush())
    return '\n'.join(result)


def strip_ug_tags(text: str) -> str:
    text = re.sub(r'\[ch\](.*?)\[/ch\]', r'\1', text)
    text = re.sub(r'\[tab\](.*?)\[/tab\]', r'\1', text, flags=re.DOTALL)
    return text


def clean_text(raw: str) -> str:
    text = normalize_whitespace(raw)
    text = strip_ug_tags(text)
    text = re.sub(r'^\s*Page\s+\d+/\d+\s*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = '\n'.join(line.rstrip() for line in text.split('\n'))
    text = re.sub(r'^\s*(\[.+?\])\s*$', r'\1', text, flags=re.MULTILINE)
    text = fix_tab_lines(text)
    return text.strip()


# ─── Scraping ─────────────────────────────────────────────────────────────────

def fetch(url: str, cookie_str: str = '') -> str:
    headers: dict[str, str] = {
        'Accept-Language': 'fr-FR,fr;q=0.9,en-US;q=0.8,en;q=0.7',
        'DNT': '1',
    }
    if cookie_str:
        headers['Cookie'] = cookie_str
    resp = requests.get(url, impersonate="chrome124", timeout=20, headers=headers)
    if resp.status_code == 403:
        print("  403 Forbidden — Cloudflare is still blocking.", file=sys.stderr)
        resp.raise_for_status()
    resp.raise_for_status()
    return resp.text


def get_ug_tab_info(html: str) -> tuple[str | None, str]:
    """
    Parse le js-store UG. Retourne (content, tab_type).
    tab_type = 'Pro' pour Guitar Pro, 'Chords'/'Tabs'/etc. pour texte.
    content = None pour les tabs GP ou si absent.
    """
    soup = BeautifulSoup(html, 'html.parser')
    div = soup.find('div', class_='js-store')
    if not div:
        return None, ''
    try:
        data = json.loads(str(div['data-content']))
        page_data = data['store']['page']['data']
        tab_type  = page_data.get('tab', {}).get('type', '')
        content   = (
            page_data.get('tab_view', {})
                     .get('wiki_tab', {})
                     .get('content', '')
        )
        return content or None, tab_type
    except (json.JSONDecodeError, KeyError, TypeError):
        return None, ''


def extract_generic(soup: BeautifulSoup) -> str:
    pres = soup.find_all('pre')
    if pres:
        return max(pres, key=lambda p: len(p.get_text())).get_text('\n')

    candidates = soup.find_all(
        ['div', 'section', 'article'],
        class_=re.compile(r'tab|chord|lyric|content|song|sheet', re.I)
    )
    if candidates:
        return max(candidates, key=lambda el: len(el.get_text())).get_text('\n')

    blocks = soup.find_all(['div', 'article', 'main', 'section'])
    if blocks:
        return max(blocks, key=lambda el: len(el.get_text())).get_text('\n')

    return soup.get_text('\n')


def scrape(url: str, dump: bool = False) -> str | None:
    """
    Retourne le contenu brut du tab (tags UG préservés pour le docx).
    Retourne None si c'est un tab Guitar Pro (géré par scrape_gp).
    """
    print(f"→ Scraping : {url}", file=sys.stderr)
    html = fetch(url)

    if dump:
        with open('ug_debug.html', 'w', encoding='utf-8') as f:
            f.write(html)
        print("  → Raw HTML saved to ug_debug.html", file=sys.stderr)

    if 'ultimate-guitar' in url:
        content, tab_type = get_ug_tab_info(html)
        if tab_type == 'Pro':
            print("  → Guitar Pro tab detected.", file=sys.stderr)
            return None
        if content:
            print("  → UG data extracted.", file=sys.stderr)
            return content
        print("  → UG JSON not found, falling back to generic extractor.", file=sys.stderr)

    soup = BeautifulSoup(html, 'html.parser')
    for tag in soup(['script', 'style', 'nav', 'footer', 'header',
                     'aside', 'iframe', 'noscript', 'button']):
        tag.decompose()

    return extract_generic(soup)


# ─── GP Download via Playwright ──────────────────────────────────────────────

_PROFILE_DIR = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), 'session', 'browser_profile'
)
def _pw_launch(p, headless: bool = False):
    os.makedirs(_PROFILE_DIR, exist_ok=True)
    last_exc: Exception = RuntimeError("No browser available")
    for channel in ('msedge', 'chrome', None):
        try:
            return p.chromium.launch_persistent_context(
                _PROFILE_DIR, channel=channel, headless=headless,
                args=[
                    '--disable-blink-features=AutomationControlled',
                    '--disable-infobars',
                    '--no-first-run',
                    '--no-default-browser-check',
                ],
                ignore_default_args=['--enable-automation'],
            )
        except Exception as e:
            last_exc = e
            continue
    raise last_exc


def scrape_gp(url: str, base_path: str, dump: bool = False) -> bool:
    """Télécharge un fichier Guitar Pro via Playwright (lien réel avec token Cloudflare)."""
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        print("  pip install playwright && playwright install", file=sys.stderr)
        return False

    with sync_playwright() as p:
        ctx = _pw_launch(p)
        page = ctx.new_page() if not ctx.pages else ctx.pages[0]
        # Masquer les signaux d'automatisation avant chaque chargement
        page.add_init_script(
            "Object.defineProperty(navigator,'webdriver',{get:()=>undefined})"
        )

        page.goto(url)
        try:
            page.wait_for_load_state('networkidle', timeout=20000)
        except Exception:
            pass

        if dump:
            with open('ug_debug.html', 'w', encoding='utf-8') as f:
                f.write(page.content())

        # Détecter si connecté via cookie bbuserid
        cookies = {c.get('name', ''): c.get('value', '') for c in ctx.cookies()}
        logged_in = cookies.get('bbuserid', '0') not in ('', '0', None)

        if not logged_in:
            print("  → Not logged in — sign in to Ultimate Guitar in the browser window.", file=sys.stderr)
            page.goto('https://www.ultimate-guitar.com/')
            input("  → Press Enter once you're signed in (keep the browser open): ")
            page.goto(url)
            try:
                page.wait_for_load_state('networkidle', timeout=20000)
            except Exception:
                pass

        # Attendre le rendu React
        try:
            page.wait_for_timeout(3000)
        except Exception:
            pass

        # Chercher le lien/bouton de téléchargement
        dl_elem = None
        for sel in [
            'a[href*="/tab/download"]',
            'a[href*="download?id"]',
            'button:has-text("Télécharger")',
            'button:has-text("Download")',
            'a:has-text("Télécharger")',
            'a:has-text("Download")',
            'button:has-text("Installer")',
        ]:
            for elem in page.query_selector_all(sel):
                try:
                    if elem.is_visible():
                        dl_elem = elem
                        break
                except Exception:
                    pass
            if dl_elem:
                break

        if not dl_elem:
            print(f"  → Download button not found (URL: {page.url})", file=sys.stderr)
            ctx.close()
            return False

        print("  → Downloading...", file=sys.stderr)

        # Écouter les downloads sur TOUTES les pages du contexte (y compris popups)
        import time as _time
        all_downloads: list = []
        page.on('download', lambda dl: all_downloads.append(dl))
        ctx.on('page', lambda new_page: new_page.on('download', lambda dl: all_downloads.append(dl)))

        dl_elem.click()
        print("  → If a CAPTCHA appears in the browser, solve it.", file=sys.stderr)

        # Attendre jusqu'à 2 minutes
        deadline = _time.time() + 120
        while _time.time() < deadline and not all_downloads:
            try:
                page.wait_for_timeout(500)
            except Exception:
                break

        if not all_downloads:
            print("  → No download detected within 2 minutes.", file=sys.stderr)
            ctx.close()
            return False

        dl = all_downloads[0]
        ext = os.path.splitext(dl.suggested_filename)[1] if dl.suggested_filename else '.gp5'
        gp_path = base_path + ext
        dl.save_as(gp_path)
        print(f"✓ {gp_path}", file=sys.stderr)
        ctx.close()
        return True


# ─── Nom de fichier depuis l'URL ─────────────────────────────────────────────

TAB_TYPES = {'chords', 'tabs', 'tab', 'bass', 'bass-tabs', 'drum', 'drums',
             'ukulele', 'power', 'pro', 'chord-pro', 'video', 'official',
             'guitar-pro'}


def filename_from_url(url: str) -> str:
    m = re.search(r'/tab/([^/]+)/([^/?#]+)', url)
    if not m:
        return 'tab.txt'

    artist_slug, song_slug = m.group(1), m.group(2)
    song_slug = re.sub(r'-\d+$', '', song_slug)
    parts = [p for p in song_slug.split('-') if p.lower() not in TAB_TYPES]
    song_slug = '-'.join(parts)

    def slugify(s: str) -> str:
        return s.replace('-', ' ').title()

    safe = re.sub(r'[<>:"/\\|?*]', '', f"{slugify(artist_slug)} - {slugify(song_slug)}")
    return f"{safe}.txt"


# ─── Export DOCX ─────────────────────────────────────────────────────────────

_CHORD_RE = re.compile(
    r'^[A-G][#b]?(?:maj|min|sus|aug|dim|add|M|m)?\d*(?:\/[A-G][#b]?)?$'
)


def _is_chord_line(line: str) -> bool:
    cleaned = re.sub(r'\*|x\d+', '', line)
    tokens = cleaned.strip().split()
    return bool(tokens) and all(_CHORD_RE.match(t) for t in tokens)


def _is_tab_line(line: str) -> bool:
    return bool(re.match(r'^\s*[eEBbGgDdAa]\|', line))


def write_docx(raw: str, path: str, title: str) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
    except ImportError:
        print("  python-docx non installé : pip install python-docx", file=sys.stderr)
        return False

    RED  = RGBColor(0xC4, 0x1E, 0x3A)
    GREY = RGBColor(0x99, 0x99, 0x99)
    FONT = 'Courier New'
    SIZE = Pt(9)

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin   = Cm(1)
        section.right_margin  = Cm(1)

    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = SIZE
    normal.paragraph_format.space_after  = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = Pt(11)

    doc.add_heading(title, level=1)

    def _para(space_before=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.line_spacing = Pt(11)
        return p

    def _run(p, text, bold=False, color=None):
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = SIZE
        if bold:
            r.bold = True
        if color:
            r.font.color.rgb = color
        return r

    text = normalize_whitespace(raw)
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'\n{3,}', '\n\n', text)

    segments = re.split(r'(\[tab\].*?\[/tab\])', text, flags=re.DOTALL)

    for segment in segments:
        if segment.startswith('[tab]') and segment.endswith('[/tab]'):
            inner = segment[5:-6]
            for line in inner.split('\n'):
                line = line.rstrip()
                stripped = line.strip()
                if not stripped:
                    continue
                if _is_tab_line(line):
                    clean = re.sub(r'\[ch\](.*?)\[/ch\]', r'\1', line)
                    _run(_para(), clean, color=GREY)
                elif '[ch]' in line:
                    p = _para()
                    for part in re.split(r'(\[ch\].*?\[/ch\])', line):
                        m = re.match(r'\[ch\](.*?)\[/ch\]', part)
                        if m:
                            _run(p, m.group(1), bold=True, color=RED)
                        else:
                            _run(p, part)
                elif _is_chord_line(stripped):
                    _run(_para(), line, bold=True, color=RED)
                else:
                    _run(_para(), line)
        else:
            blank_count = 0
            for line in segment.split('\n'):
                line = line.rstrip()
                stripped = line.strip()

                if not stripped:
                    blank_count += 1
                    if blank_count == 1:
                        _para()
                    continue
                blank_count = 0

                if re.match(r'^\[[^\[\]/]+\]$', stripped):
                    _run(_para(space_before=5), stripped, bold=True)
                    continue

                if '[ch]' in line:
                    p = _para()
                    for part in re.split(r'(\[ch\].*?\[/ch\])', line):
                        m = re.match(r'\[ch\](.*?)\[/ch\]', part)
                        if m:
                            _run(p, m.group(1), bold=True, color=RED)
                        else:
                            _run(p, part)
                    continue

                if _is_chord_line(stripped):
                    _run(_para(), line, bold=True, color=RED)
                    continue

                _run(_para(), line)

    doc.save(path)
    return True


# ─── Mode nettoyage (stdin) ───────────────────────────────────────────────────

def clean_stdin() -> str:
    print("Paste the text, then Ctrl+Z + Enter:", file=sys.stderr)
    return clean_text(sys.stdin.read())


# ─── CLI ──────────────────────────────────────────────────────────────────────

OUTPUT_DIR = 'output'


def main():
    parser = argparse.ArgumentParser(
        description='Scrape and clean guitar tabs.',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemples :
  python app.py "https://tabs.ultimate-guitar.com/tab/neil-young/natural-beauty-chords-88512"
  python app.py "https://tabs.ultimate-guitar.com/tab/iron-maiden/2-minutes-to-midnight-guitar-pro-222022"
  python app.py --clean
"""
    )
    parser.add_argument('url', nargs='?', help='Tab URL (wrap in quotes)')
    parser.add_argument('--clean', action='store_true',
                        help='Clean mode: read from stdin')
    parser.add_argument('--dump', action='store_true',
                        help='Save raw HTML to ug_debug.html (debug)')
    args = parser.parse_args()

    if args.clean:
        print(clean_stdin())
        return

    if not args.url:
        parser.print_help()
        sys.exit(1)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    base      = re.sub(r'\.txt$', '', filename_from_url(args.url))
    base_path = os.path.join(OUTPUT_DIR, base)

    raw = scrape(args.url, dump=args.dump)

    if raw is None:
        # Tab Guitar Pro
        scrape_gp(args.url, base_path, dump=args.dump)
        return

    txt_path = base_path + '.txt'
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(clean_text(raw))
    print(f"✓ {txt_path}", file=sys.stderr)

    docx_path = base_path + '.docx'
    if write_docx(raw, docx_path, base):
        print(f"✓ {docx_path}", file=sys.stderr)


if __name__ == '__main__':
    main()
